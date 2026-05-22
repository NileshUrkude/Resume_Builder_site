import io
import re
from collections import Counter

from django.contrib.auth.hashers import check_password, make_password
from django.db import transaction
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.utils import timezone
from xhtml2pdf import pisa

from .models import (
    Achievement,
    Certificate,
    Education,
    Experience,
    Hobby,
    Language,
    Project,
    Resume,
    ResumeEvent,
    Skill,
)

TEMPLATE_MAP = {
    't1': 'resumes/t1.html',
    't2': 'resumes/t2.html',
    't3': 'resumes/t3.html',
    't4': 'resumes/t4.html',
    't5': 'resumes/t5.html',
    't6': 'resumes/t6.html',
    't7': 'resumes/t7.html',
    't8': 'resumes/t8.html',
}

VALID_TEMPLATES = frozenset(TEMPLATE_MAP.keys())


def get_user_resume(request, resume_id=None):
    qs = Resume.objects.filter(user=request.user)
    if resume_id:
        return qs.filter(pk=resume_id).first()
    active = request.session.get('active_resume_id')
    if active:
        resume = qs.filter(pk=active).first()
        if resume:
            return resume
    return qs.first()


def set_active_resume(request, resume):
    request.session['active_resume_id'] = resume.pk


def track_event(resume, event_type, template=''):
    if resume:
        ResumeEvent.objects.create(resume=resume, event_type=event_type, template=template)


def set_share_password(resume, raw_password):
    if raw_password:
        resume.share_password = make_password(raw_password)
    else:
        resume.share_password = ''


def check_share_password(resume, raw_password):
    if not resume.share_password:
        return True
    return check_password(raw_password, resume.share_password)


@transaction.atomic
def duplicate_resume(source: Resume, new_name=None):
    clone = Resume.objects.create(
        user=source.user,
        resume_name=new_name or f'{source.resume_name} (Copy)',
        is_public=False,
        preferred_template=source.preferred_template,
        accent_color=source.accent_color,
        font_family=source.font_family,
        title=source.title,
        full_name=source.full_name,
        email=source.email,
        phone=source.phone,
        summary=source.summary,
        github=source.github,
        linkedin=source.linkedin,
    )
    if source.photo:
        clone.photo = source.photo
        clone.save(update_fields=['photo'])

    for edu in source.educations.all():
        Education.objects.create(
            resume=clone, degree=edu.degree, college=edu.college, location=edu.location,
            start_date=edu.start_date, end_date=edu.end_date, duration=edu.duration, order=edu.order,
        )
    for exp in source.experiences.all():
        Experience.objects.create(
            resume=clone, job_title=exp.job_title, company=exp.company,
            start_date=exp.start_date, end_date=exp.end_date, duration=exp.duration,
            description=exp.description, order=exp.order,
        )
    for proj in source.projects.all():
        Project.objects.create(
            resume=clone, name=proj.name, description=proj.description,
            tech_stack=proj.tech_stack, order=proj.order,
        )
    for sk in source.skills.all():
        Skill.objects.create(resume=clone, name=sk.name, level=sk.level, order=sk.order)
    for cert in source.certificates.all():
        Certificate.objects.create(resume=clone, name=cert.name, order=cert.order)
    for ach in source.achievements.all():
        Achievement.objects.create(resume=clone, title=ach.title, order=ach.order)
    for lang in source.languages.all():
        Language.objects.create(resume=clone, name=lang.name, proficiency=lang.proficiency, order=lang.order)
    for hobby in source.hobbies.all():
        Hobby.objects.create(resume=clone, name=hobby.name, order=hobby.order)
    return clone


def generate_summary(resume: Resume) -> str:
    """Rule-based professional summary (no external API)."""
    roles = [e.job_title for e in resume.experiences.all()[:3]]
    skills = [s.name for s in resume.skills.all()[:8]]
    parts = []
    if roles:
        parts.append(f'Experienced professional with background as {", ".join(roles)}.')
    if resume.summary:
        parts.append(resume.summary.strip())
    elif resume.projects.exists():
        parts.append(f'Built projects including {resume.projects.first().name}.')
    if skills:
        parts.append(f'Skilled in {", ".join(skills)}.')
    return ' '.join(parts)[:600] or 'Motivated professional seeking new opportunities.'


def compute_ats_score(resume: Resume) -> dict:
    """Simple ATS checklist scoring."""
    checks = []
    score = 0
    max_score = 100

    if resume.full_name and len(resume.full_name) >= 3:
        checks.append({'ok': True, 'text': 'Full name present'})
        score += 10
    else:
        checks.append({'ok': False, 'text': 'Add your full name'})

    if resume.email and '@' in resume.email:
        checks.append({'ok': True, 'text': 'Email present'})
        score += 10
    else:
        checks.append({'ok': False, 'text': 'Add a professional email'})

    if resume.phone:
        checks.append({'ok': True, 'text': 'Phone number present'})
        score += 10
    else:
        checks.append({'ok': False, 'text': 'Add phone number'})

    summary_len = len(resume.summary or '')
    if 80 <= summary_len <= 600:
        checks.append({'ok': True, 'text': 'Summary length is ATS-friendly (80–600 chars)'})
        score += 15
    elif summary_len > 0:
        checks.append({'ok': False, 'text': f'Summary is {summary_len} chars — aim for 80–600'})
        score += 5
    else:
        checks.append({'ok': False, 'text': 'Add a professional summary'})

    if resume.experiences.exists():
        checks.append({'ok': True, 'text': 'Experience section filled'})
        score += 20
    else:
        checks.append({'ok': False, 'text': 'Add at least one experience'})

    if resume.skills.count() >= 5:
        checks.append({'ok': True, 'text': f'{resume.skills.count()} skills listed (good)'})
        score += 15
    elif resume.skills.exists():
        checks.append({'ok': False, 'text': 'Add more skills (target 5+)'})
        score += 8
    else:
        checks.append({'ok': False, 'text': 'Add skills section'})

    if resume.educations.exists():
        checks.append({'ok': True, 'text': 'Education section present'})
        score += 10
    else:
        checks.append({'ok': False, 'text': 'Add education'})

    if resume.linkedin or resume.github:
        checks.append({'ok': True, 'text': 'Online profile links included'})
        score += 10
    else:
        checks.append({'ok': False, 'text': 'Add LinkedIn or GitHub URL'})

    word_count = len((resume.summary or '').split())
    for exp in resume.experiences.all():
        word_count += len(exp.description.split())
    if word_count >= 150:
        checks.append({'ok': True, 'text': 'Enough content for ATS parsing'})
        score += 10
    else:
        checks.append({'ok': False, 'text': 'Expand descriptions (150+ words total)'})

    return {
        'score': min(score, max_score),
        'max_score': max_score,
        'checks': checks,
        'grade': 'Excellent' if score >= 85 else 'Good' if score >= 70 else 'Needs work',
    }


def match_job_description(resume: Resume, jd_text: str) -> dict:
    jd_words = set(re.findall(r'[a-zA-Z+#]{2,}', jd_text.lower()))
    skill_names = {s.name.lower() for s in resume.skills.all()}
    resume_words = skill_names.copy()
    for exp in resume.experiences.all():
        resume_words.update(re.findall(r'[a-zA-Z+#]{2,}', exp.description.lower()))
        resume_words.update(re.findall(r'[a-zA-Z+#]{2,}', exp.job_title.lower()))
    for proj in resume.projects.all():
        resume_words.update(re.findall(r'[a-zA-Z+#]{2,}', proj.tech_stack.lower()))

    common_tech = {
        'python', 'django', 'javascript', 'react', 'java', 'sql', 'html', 'css',
        'node', 'aws', 'docker', 'kubernetes', 'git', 'api', 'rest', 'mongodb',
        'postgresql', 'excel', 'communication', 'leadership', 'agile', 'scrum',
    }
    jd_keywords = [w for w in jd_words if w in common_tech or len(w) > 4]
    jd_keywords = list(dict.fromkeys(jd_keywords))[:40]

    matched = [k for k in jd_keywords if k in resume_words]
    missing = [k for k in jd_keywords if k not in resume_words][:20]

    pct = int(len(matched) / len(jd_keywords) * 100) if jd_keywords else 0
    return {
        'match_percent': pct,
        'matched': matched,
        'missing': missing,
        'suggestions': missing[:8],
    }


def render_resume_html(resume, template, for_pdf=False):
    accent = resume.accent_color or '#4f46e5'
    font = resume.font_family or 'Arial'
    return render_to_string(
        TEMPLATE_MAP.get(template, TEMPLATE_MAP['t1']),
        {'resume': resume, 'for_pdf': for_pdf, 'accent_color': accent, 'font_family': font},
    )


def pdf_response(resume, template, filename='resume.pdf'):
    html = render_resume_html(resume, template, for_pdf=True)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    pisa.CreatePDF(html, dest=response)
    return response


def export_docx(resume, template='t1'):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    accent = resume.accent_color.lstrip('#')
    try:
        r = int(accent[0:2], 16)
        g = int(accent[2:4], 16)
        b = int(accent[4:6], 16)
    except (ValueError, IndexError):
        r, g, b = 79, 70, 229

    title = doc.add_heading(resume.full_name, 0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(r, g, b)

    doc.add_paragraph(resume.title)
    doc.add_paragraph(f'{resume.email} | {resume.phone}')
    if resume.linkedin:
        doc.add_paragraph(f'LinkedIn: {resume.linkedin}')
    if resume.github:
        doc.add_paragraph(f'GitHub: {resume.github}')

    if resume.summary:
        doc.add_heading('Summary', level=2)
        doc.add_paragraph(resume.summary)

    if resume.experiences.exists():
        doc.add_heading('Experience', level=2)
        for exp in resume.experiences.all():
            p = doc.add_paragraph()
            p.add_run(f'{exp.job_title} — {exp.company}').bold = True
            doc.add_paragraph(exp.date_range())
            for bullet in exp.bullet_points() or [exp.description]:
                doc.add_paragraph(bullet, style='List Bullet')

    if resume.educations.exists():
        doc.add_heading('Education', level=2)
        for edu in resume.educations.all():
            doc.add_paragraph(f'{edu.degree} — {edu.college} ({edu.date_range()})')

    if resume.projects.exists():
        doc.add_heading('Projects', level=2)
        for proj in resume.projects.all():
            doc.add_paragraph(f'{proj.name} ({proj.tech_stack})')
            doc.add_paragraph(proj.description)

    if resume.skills.exists():
        doc.add_heading('Skills', level=2)
        skills_text = ', '.join(f'{s.name} ({s.get_level_display()})' for s in resume.skills.all())
        doc.add_paragraph(skills_text)

    if resume.languages.exists():
        doc.add_heading('Languages', level=2)
        doc.add_paragraph(', '.join(f'{l.name} ({l.get_proficiency_display()})' for l in resume.languages.all()))

    if resume.hobbies.exists():
        doc.add_heading('Hobbies', level=2)
        doc.add_paragraph(', '.join(h.name for h in resume.hobbies.all()))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{resume.resume_name.replace(" ", "_")}.docx"'
    return response


def qr_png_response(url: str):
    import qrcode
    img = qrcode.make(url)
    buffer = io.BytesIO()
    img.save(buffer, format='PNG')
    buffer.seek(0)
    return HttpResponse(buffer.read(), content_type='image/png')


def save_ordered_formset(formset, resume, related_name, order_prefix):
    getattr(resume, related_name).all().delete()
    for idx, form in enumerate(formset):
        if form.cleaned_data and not form.cleaned_data.get('DELETE'):
            obj = form.save(commit=False)
            obj.resume = resume
            obj.order = idx
            obj.save()
